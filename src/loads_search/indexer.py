"""
Whoosh index: build and rebuild from file metadata + file contents.
Index is local per device; rebuild from file_metadata.json when needed.
"""
import shutil
from pathlib import Path
from typing import Any

from whoosh import index
from whoosh.fields import ID, KEYWORD, TEXT, Schema
from whoosh.qparser import QueryParser

from .paths import get_logs_dir, get_search_index_dir
from .metadata import load_metadata, save_metadata

# Schema: path = unique id (file path or "command:..."), result_type = "file" | "command", content = searchable text
SCHEMA = Schema(
    path=ID(stored=True, unique=True),
    result_type=KEYWORD(stored=True),
    content=TEXT(stored=True),
)


def _read_content(path_str: str, max_chars: int = 500_000) -> str:
    """Read file as text; return empty string on error or if too large. Handles .docx and .pdf files."""
    path = Path(path_str)
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    
    # PDF documents: extract text with pdfplumber (primary) or PyPDF2 (fallback)
    if suffix == ".pdf":
        # Try pdfplumber first (better encoding support)
        try:
            import pdfplumber
            import threading
            import queue
            
            def extract_with_timeout():
                """Extract PDF text with timeout protection"""
                text_parts = []
                try:
                    with pdfplumber.open(path_str) as pdf:
                        # Limit to first 50 pages to prevent hanging on large e-books
                        max_pages = min(50, len(pdf.pages))
                        for i, page in enumerate(pdf.pages[:max_pages]):
                            try:
                                text = page.extract_text()
                                if text and text.strip():
                                    text_parts.append(text)
                                    # Stop if we've collected enough content
                                    if len("\n".join(text_parts)) > max_chars:
                                        break
                            except Exception:
                                continue
                except Exception as e:
                    raise e
                return "\n".join(text_parts) if text_parts else ""
            
            # Use threading timeout (works on Windows)
            result_queue = queue.Queue()
            def worker():
                try:
                    result = extract_with_timeout()
                    result_queue.put(("success", result))
                except Exception as e:
                    result_queue.put(("error", str(e)))
            
            thread = threading.Thread(target=worker)
            thread.daemon = True
            thread.start()
            thread.join(timeout=30)  # 30 second timeout
            
            if thread.is_alive():
                # Thread is still running - likely stuck on a large PDF
                print(f"⚠️ PDF extraction timeout for {path_str} - file too large")
                return ""
            
            if not result_queue.empty():
                status, result = result_queue.get()
                if status == "success" and result:
                    return result[:max_chars] if len(result) > max_chars else result
                    
        except ImportError:
            pass
        except Exception:
            pass
        
        # Fallback to PyPDF2 with similar optimizations
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path_str)
            text_parts = []
            
            # Limit to first 50 pages
            max_pages = min(50, len(reader.pages))
            for i in range(max_pages):
                try:
                    page = reader.pages[i]
                    text = page.extract_text()
                    if text and text.strip():
                        # Try to fix common Turkish character encoding issues
                        text = text.replace('■', 'ş').replace('■', 'Ş').replace('■', 'ç').replace('■', 'Ç')
                        text = text.replace('■', 'ğ').replace('■', 'Ğ').replace('■', 'ı').replace('■', 'İ')
                        text = text.replace('■', 'ö').replace('■', 'Ö').replace('■', 'ü').replace('■', 'Ü')
                        text_parts.append(text)
                        
                        # Stop if we've collected enough content
                        if len("\n".join(text_parts)) > max_chars:
                            break
                except Exception:
                    continue
            text = "\n".join(text_parts)
            return text[:max_chars] if len(text) > max_chars else text
        except Exception:
            return ""
    
    # Word documents: extract text with python-docx
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(path_str)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
            return text[:max_chars] if len(text) > max_chars else text
        except Exception:
            return ""
    
    # Plain text - use chardet for encoding detection with Turkish fallbacks
    try:
        size = path.stat().st_size
        if size > max_chars * 2:  # rough: 2 bytes per char
            return ""
        raw = path.read_bytes()
        
        # First try chardet for automatic detection
        try:
            import chardet
            detected = chardet.detect(raw)
            if detected and detected['confidence'] > 0.7:
                encoding = detected['encoding']
                if encoding:
                    try:
                        return raw.decode(encoding)
                    except UnicodeDecodeError:
                        pass
        except ImportError:
            pass
        
        # Fallback to common encodings with Turkish support
        encodings = ['utf-8', 'utf-8-sig', 'windows-1254', 'iso-8859-9', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Final fallback to utf-8 with error replacement
        return raw.decode("utf-8", errors="replace")
    except (OSError, UnicodeDecodeError):
        return ""


def build_index(entries: list[dict[str, Any]], command_entries: list[dict[str, Any]] | None = None) -> int:
    """
    Build Whoosh index from file entries and optional command entries.
    Files: list of {path, mtime, size}. Commands: list of {type, timestamp, command, cwd}.
    Returns total number of documents indexed.
    """
    idx_dir = get_search_index_dir()
    idx_dir.mkdir(parents=True, exist_ok=True)
    if idx_dir.exists() and any(idx_dir.iterdir()):
        shutil.rmtree(idx_dir)
    idx_dir.mkdir(parents=True, exist_ok=True)

    ix = index.create_in(str(idx_dir), SCHEMA)
    writer = ix.writer()
    count = 0
    for e in entries:
        path_str = e.get("path") or ""
        if not path_str:
            continue
        content = _read_content(path_str)
        try:
            writer.add_document(path=path_str, result_type="file", content=content or " ")
            count += 1
        except Exception:
            continue
    # Index commands from activity logs
    from .activity_logger import load_commands_from_logs
    cmd_list = command_entries if command_entries is not None else load_commands_from_logs(get_logs_dir())
    for i, c in enumerate(cmd_list):
        cmd = (c.get("command") or "").strip()
        if not cmd:
            continue
        ts = c.get("timestamp", "") or str(i)
        doc_id = f"command:{ts}:{i}"
        try:
            writer.add_document(path=doc_id, result_type="command", content=cmd)
            count += 1
        except Exception:
            continue
    writer.commit()
    return count


def rebuild_index() -> int:
    """Load file_metadata.json and build index. Returns number of documents indexed."""
    entries = load_metadata()
    return build_index(entries)


def full_index(config: dict[str, Any]) -> int:
    """
    Sync terminal history (if enabled), crawl folders, save metadata, build Whoosh index (files + commands).
    Use this for "Re-index" or first-time index.
    """
    from .crawler import crawl
    from .activity_logger import sync_terminal_history

    if config.get("log_terminal_history", True):
        get_logs_dir().mkdir(parents=True, exist_ok=True)
        sync_terminal_history(get_logs_dir())

    entries = list(crawl(
        folders=config.get("folders_to_index", []),
        exclude_patterns=config.get("exclude_patterns", []),
        max_file_size_kb=int(config.get("max_file_size_kb", 512)),
    ))
    save_metadata(entries)
    return build_index(entries)


def get_index():
    """Open existing index for reading. Returns None if index does not exist or is empty."""
    idx_dir = get_search_index_dir()
    if not idx_dir.exists():
        return None
    try:
        if not index.exists_in(str(idx_dir)):
            return None
        return index.open_dir(str(idx_dir))
    except Exception:
        return None


def search_index(q: str, limit: int = 50) -> list[tuple[str, str, str, str | None]]:
    """
    Query Whoosh. Returns list of (path_or_id, snippet, result_type, copy_text).
    result_type is "file" or "command". copy_text is set for commands (for clipboard); None for files.
    """
    q = (q or "").strip()
    if not q:
        return []
    ix = get_index()
    if ix is None:
        return []
    with ix.searcher() as searcher:
        parser = QueryParser("content", schema=SCHEMA)
        try:
            query = parser.parse(q)
        except Exception:
            return []
        results = searcher.search(query, limit=limit)
        out = []
        for hit in results:
            path = hit.get("path") or ""
            result_type = hit.get("result_type") or "file"
            content = hit.get("content") or ""
            snippet = content[:200].replace("\n", " ").strip()
            if len(content) > 200:
                snippet += "..."
            copy_text = content if result_type == "command" else None
            out.append((path, snippet, result_type, copy_text))
        return out
